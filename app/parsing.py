"""スキルシートファイル（xlsx / docx）からプレーンテキストを抽出する。"""
from __future__ import annotations

import io

from docx import Document
from openpyxl import load_workbook


def extract_text(filename: str, content: bytes) -> str:
    lower = filename.lower()
    if lower.endswith(".xlsx"):
        return _extract_xlsx(content)
    if lower.endswith(".docx"):
        return _extract_docx(content)
    if lower.endswith(".txt") or lower.endswith(".md"):
        return content.decode("utf-8", errors="ignore")
    raise ValueError(f"対応していないファイル形式です: {filename}")


def _extract_xlsx(content: bytes) -> str:
    workbook = load_workbook(io.BytesIO(content), data_only=True)
    lines: list[str] = []
    for sheet in workbook.worksheets:
        lines.append(f"# シート: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None and str(c).strip() != ""]
            if cells:
                lines.append("\t".join(cells))
    return "\n".join(lines)


def _extract_docx(content: bytes) -> str:
    document = Document(io.BytesIO(content))
    lines: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            lines.append(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append("\t".join(cells))
    return "\n".join(lines)
